"""
Import real university proposals from dataset/university_proposals/ into project_ideas.

Run from backend/:
  python -m scripts.import_university_proposals
  python -m scripts.import_university_proposals --dry-run
  python -m scripts.import_university_proposals --generate-sql

Does NOT run relevancy analysis or modify AI/auth/notification code.
"""

from __future__ import annotations

import argparse
import asyncio
import importlib.util
import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sqlalchemy import select, text

_spec = importlib.util.spec_from_file_location(
    "bootstrap", Path(__file__).resolve().parent / "_bootstrap.py"
)
_bootstrap = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_bootstrap)

from app.database import AsyncSessionLocal
from app.models.professor import Professor
from app.models.project import ProjectIdea, ProjectStatus
from app.models.user import User

ROOT = Path(__file__).resolve().parents[2]
PROPOSALS_DIR = ROOT / "dataset" / "university_proposals"
AUDIT_JSON = ROOT / "dataset" / "university_proposals_audit.json"
SQL_OUT = ROOT / "university_proposals_import.sql"
REPORT_JSON = ROOT / "dataset" / "university_proposals_import_result.json"

PROFESSOR_EMAIL = "professor@uol.edu.pk"

SKIP_FILENAMES = {
    "FYP Proposal.pdf",
    "Proposal.pdf",
    "pixelwave (hrm) project proposal by group 2_260621_125121.pdf",
}

SECTION_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("problem_statement", re.compile(r"problem\s*(?:statement|identification)", re.I)),
    ("proposed_solution", re.compile(r"(?:proposed\s*solution|problem\s*solution|methodology|approach)", re.I)),
    ("technologies", re.compile(r"(?:tools?\s*(?:&|and)\s*technologies?|technologies?|tech\s*stack)", re.I)),
    ("target_users", re.compile(r"(?:target\s*users?|target\s*audience|primary\s*users?)", re.I)),
    ("expected_impact", re.compile(r"(?:expected\s*(?:output|impact|outcome)|significance\s*of)", re.I)),
    ("unique_features", re.compile(r"(?:unique\s*features?|main\s*features?|key\s*features?)", re.I)),
    ("innovation_aspect", re.compile(r"(?:innovation|novelty|contribution)", re.I)),
    ("description", re.compile(r"(?:introduction|abstract|project\s*overview|background)", re.I)),
]

# Curated overrides from UNIVERSITY_PROPOSAL_AUDIT.md (clean titles + gap fills)
CURATED: dict[str, dict[str, str]] = {
    "ahad CipherPlay_Proposal.pdf": {
        "title": "CipherPlay — AI-Powered Children's Indoor Fitness Game",
        "description": (
            "CipherPlay is a mobile game that combines real-time YOLO object detection with "
            "physical movement challenges for children aged 4–14, turning screen time into "
            "indoor physical activity under parental supervision."
        ),
        "problem_statement": (
            "Children aged 4–14 average high daily passive screen time, leading to sedentary "
            "behaviour and reduced physical activity. Existing mobile games provide entertainment "
            "but fail to bridge screen time with physical engagement indoors."
        ),
        "proposed_solution": (
            "Build an Android application using on-device YOLO object detection, adaptive "
            "gameplay tied to detected objects, and physical entropy-based encryption for "
            "secure in-game rewards."
        ),
        "technologies": "Python, YOLO, Android, PyTorch, OpenCV",
        "target_users": "Children aged 4–14 and parents seeking active screen-time alternatives",
        "expected_impact": (
            "A functional Android app combining real-time AI object detection, adaptive gameplay, "
            "and hardware-backed encryption to promote physical activity."
        ),
        "unique_features": (
            "On-device YOLO detection; physical entropy encryption; indoor-safe gameplay without GPS"
        ),
        "innovation_aspect": (
            "Applies YOLO object detection and true random encryption to children's indoor fitness "
            "gaming — a combination not present in mainstream AR outdoor games."
        ),
        "category": "Gaming & Entertainment",
        "target_industry": "Entertainment",
    },
    "AI_Classroom_Noise_Notes_Proposal.pdf": {
        "title": "AI-Powered Real-Time Classroom Noise Detection and Automatic Lecture Notes Generator",
        "description": (
            "An AI classroom assistant that classifies noise in real time, filters audio, "
            "transcribes teacher speech with Whisper, and generates summarized lecture notes."
        ),
        "problem_statement": (
            "Traditional lecture recordings contain background noise, poor audio quality, and lack "
            "structured summaries, forcing students to re-listen to long recordings."
        ),
        "proposed_solution": (
            "Pipeline: noise classification (CNN/CRNN) → real-time filtering → Whisper ASR → "
            "LLM summarization, delivered via Python backend and web UI."
        ),
        "technologies": "Python, PyTorch, Whisper, Transformers, Librosa, Streamlit, FastAPI",
        "target_users": "University students and lecturers",
        "expected_impact": "Clean audio recordings, searchable transcripts, and auto-generated lecture notes",
        "unique_features": "Real-time noise classification; combined filtering, transcription, and summarization",
        "innovation_aspect": "Integrates audio AI, speech recognition, and NLP summarization in one classroom tool",
        "category": "Education & EdTech",
        "target_industry": "Higher Education",
    },
    "Filled_FYP_Proposal_AI_Tuition_System.docx": {
        "title": "AI Smart Tuition Recommendation System",
        "target_users": "Students, parents, and private tutors seeking matched tuition arrangements",
        "expected_impact": (
            "Improved tutor-student matching, better academic outcomes, and reduced wasted tuition spending"
        ),
        "technologies": "React, Flutter, Node.js, Django, Python, Scikit-learn",
        "unique_features": "AI tutor matching; weak-subject analysis; personalized study planner",
        "innovation_aspect": "ML-driven tutor recommendation using performance, budget, and location constraints",
        "category": "Education & EdTech",
        "target_industry": "Higher Education",
    },
    "Final year project.pdf": {
        "title": "ScholarIQ — AI Scholarship Recommendation Platform",
        "proposed_solution": (
            "An AI-powered web platform that recommends genuine international scholarships based on "
            "academic profile, nationality, and goals, with deadline alerts and fraud filtering."
        ),
        "technologies": "React.js, Bootstrap 5, Python, FastAPI, PostgreSQL",
        "target_users": "International students seeking scholarships",
        "unique_features": "Profile-based scholarship matching; deadline alerts; fraud filtering",
        "innovation_aspect": "Centralized intelligent scholarship discovery instead of scattered manual search",
        "category": "Education & EdTech",
        "target_industry": "Higher Education",
    },
    "FYP Pre.docx": {
        "title": "Smart Relocation Assistance and Property Finder System",
        "technologies": "React.js, Node.js, PostgreSQL, Google Maps API",
        "target_users": "Individuals relocating to new cities in Pakistan for residence or investment",
        "unique_features": "Structured property search; dealer ratings; integrated relocation services",
        "innovation_aspect": "Combines property discovery with relocation logistics in one platform",
        "category": "General Software / AI",
        "target_industry": "Information Technology",
    },
    "FYP Proposal.docx": {
        "title": "AI-Based Retinal Disease Detection System Using Deep Learning",
        "description": (
            "A deep learning system to detect diabetic retinopathy and glaucoma from retinal fundus "
            "images, supporting early diagnosis especially in underserved areas."
        ),
        "problem_statement": (
            "Manual retinal diagnosis is time-consuming and inaccessible in rural areas, delaying "
            "treatment for diabetic retinopathy and glaucoma."
        ),
        "target_users": "Patients, ophthalmologists, and rural clinic staff",
        "technologies": "Python, TensorFlow, Keras, OpenCV, React, FastAPI",
        "unique_features": "Fundus image upload; CNN classification; remote monitoring support",
        "innovation_aspect": "Automated CNN-based retinal screening accessible via a web application",
        "category": "Healthcare & Medical AI",
        "target_industry": "Healthcare",
    },
    "FYP_C-PRMS 2.pdf": {
        "title": "Integrated Health Information System (IHIS)",
        "proposed_solution": (
            "Develop an interoperable health information system with patient registration, EHR modules, "
            "and Scrum-based iterative delivery for Pakistani healthcare institutions."
        ),
        "target_users": "Healthcare institutions, patients, and government health planners",
        "technologies": "Jira, Git, AWS/Azure, Docker, Python, PostgreSQL",
        "unique_features": "Patient registration; EHR interoperability; sprint-based delivery",
        "innovation_aspect": "National-scale health information integration addressing paper-record silos",
        "category": "Healthcare & Medical AI",
        "target_industry": "Healthcare",
    },
    "FYP_Proposal_1.docx": {
        "title": 'AI-Based "Can I Resell This?" System for Waste-to-Value Optimization',
        "expected_impact": (
            "Reduce waste and financial loss by guiding users to sell, reuse, repair, or recycle items "
            "based on AI condition and value analysis"
        ),
        "technologies": "React, Flutter, FastAPI, TensorFlow, PyTorch, OpenCV, Firebase",
        "target_users": "General consumers and sustainability-conscious households",
        "unique_features": "Image-based object detection; condition analysis; resale price prediction",
        "innovation_aspect": "Combines YOLO detection with resale/recycle decision support in one workflow",
        "category": "General Software / AI",
        "target_industry": "Environment & Sustainability",
    },
    "Proposal CareerCraft V2.pdf": {
        "title": "CareerCraft AI — Intelligent Interview and Resume Mastery Platform",
        "description": (
            "CareerCraft AI provides AI interview simulation, resume ATS optimization, and "
            "communication feedback for job seekers."
        ),
        "problem_statement": (
            "Job seekers lack realistic interview practice, personalized resume feedback, and "
            "communication coaching beyond generic online resources."
        ),
        "proposed_solution": (
            "Build an AI platform with adaptive mock interviews, resume optimization, skill-gap "
            "detection, and progress dashboards for students and professionals."
        ),
        "technologies": "Python, TensorFlow, PyTorch, OpenCV, spaCy, NLTK, Flutter, Node.js",
        "target_users": "University students, graduates, and career-switching professionals",
        "unique_features": "AI mock interviews; ATS resume scoring; verbal and non-verbal feedback",
        "innovation_aspect": "Holistic career readiness combining NLP interview analysis and resume optimization",
        "category": "Career & Employment",
        "target_industry": "Career Services",
    },
    "qaiser11 fyp perposol.pdf": {
        "title": "AI-Powered Career Coaching and Job Preparation Platform",
        "description": (
            "A unified career platform for resume building, cover letter generation, and AI mock "
            "interviews using modern web and voice AI stack."
        ),
        "problem_statement": (
            "Job preparation is fragmented across separate tools for resumes, cover letters, and "
            "interviews, producing generic results for candidates."
        ),
        "proposed_solution": (
            "Integrate resume rewriting, tailored cover letters, dynamic interview question banks, "
            "and voice-based mock interviews in a single Next.js application."
        ),
        "technologies": "Next.js, Gemini API, WebRTC, Vapi, Clerk, Prisma, PostgreSQL",
        "target_users": "Students, fresh graduates, and career switchers",
        "expected_impact": (
            "Higher-quality applications and interview readiness through personalized AI coaching"
        ),
        "unique_features": "Resume builder; cover letter generator; voice mock interviews",
        "innovation_aspect": "Single platform replacing fragmented career prep tools with AI personalization",
        "category": "Career & Employment",
        "target_industry": "Career Services",
    },
}


@dataclass
class ImportResult:
    filename: str
    status: str
    title: str = ""
    project_id: int | None = None
    reason: str = ""


def _clean(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(p.extract_text() or "" for p in reader.pages)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_section(text: str, pattern: re.Pattern[str]) -> str:
    m = pattern.search(text)
    if not m:
        return ""
    start = m.end()
    next_header = re.search(
        r"\n\s*(?:\d+[\.\)]\s+)?[A-Z][A-Za-z\s/&]{3,40}\s*[:\-]",
        text[start : start + 2500],
    )
    end = start + (next_header.start() if next_header else 800)
    return _clean(text[start:end])[:2000]


def parse_from_text(text: str, filename: str) -> dict[str, str]:
    data: dict[str, str] = {}
    for key, pattern in SECTION_PATTERNS:
        val = extract_section(text, pattern)
        if val:
            data[key] = val
    if not data.get("title"):
        stem = Path(filename).stem.replace("_", " ").replace("-", " ")
        data["title"] = _clean(stem).title()
    if not data.get("description") and text:
        intro = extract_section(text, re.compile(r"introduction", re.I))
        if intro:
            data["description"] = intro
        elif len(text) > 100:
            data["description"] = _clean(text[:600])
    if not data.get("technologies"):
        m = re.search(
            r"(Python|React|Flutter|TensorFlow|Node\.js)[^\n]{5,200}",
            text,
            re.I,
        )
        if m:
            data["technologies"] = _clean(m.group(0))[:400]
    return data


def build_proposal_record(filename: str) -> dict[str, str]:
    path = PROPOSALS_DIR / filename
    raw = extract_docx(path) if path.suffix.lower() == ".docx" else extract_pdf(path)
    parsed = parse_from_text(raw, filename)
    curated = CURATED.get(filename, {})
    merged = {**parsed, **curated}

    required_defaults = {
        "description": merged.get("description") or _clean(raw[:500]),
        "technologies": merged.get("technologies") or "Python, Web Application",
        "problem_statement": merged.get("problem_statement") or "",
        "proposed_solution": merged.get("proposed_solution") or "",
        "unique_features": merged.get("unique_features") or "",
        "innovation_aspect": merged.get("innovation_aspect") or "",
        "target_users": merged.get("target_users") or "",
        "expected_impact": merged.get("expected_impact") or "",
        "category": merged.get("category") or "General Software / AI",
        "target_industry": merged.get("target_industry") or "Information Technology",
    }
    merged.update({k: v for k, v in required_defaults.items() if not merged.get(k)})

    title = merged.get("title") or Path(filename).stem
    title = _clean(title)[:500]
    merged["title"] = title
    merged["description"] = _clean(merged["description"])[:5000]
    merged["technologies"] = _clean(merged["technologies"])[:500]
    for key in (
        "problem_statement",
        "proposed_solution",
        "unique_features",
        "innovation_aspect",
        "target_users",
        "expected_impact",
    ):
        merged[key] = _clean(merged.get(key, ""))[:3000]
    merged["category"] = _clean(merged["category"])[:255]
    merged["target_industry"] = _clean(merged["target_industry"])[:255]
    merged["_source_file"] = filename
    return merged


def esc_sql(value: str) -> str:
    return value.replace("'", "''")


def generate_sql(proposals: list[dict[str, str]]) -> str:
    lines = [
        "-- University proposals import (10 ready proposals)",
        "-- Run: psql -U postgres -d fyp_relevancy_system -f university_proposals_import.sql",
        "-- Prefer: python -m scripts.import_university_proposals (handles escaping safely)",
        "BEGIN;",
        "",
    ]
    for i, p in enumerate(proposals):
        offset = i % 3
        lines.append(f"-- Source: {p['_source_file']}")
        lines.append("INSERT INTO project_ideas (")
        lines.append(
            "  student_id, professor_id, title, technologies, description,"
        )
        lines.append(
            "  category, target_industry, problem_statement, proposed_solution,"
        )
        lines.append(
            "  unique_features, innovation_aspect, target_users, expected_impact,"
        )
        lines.append("  professor_email, status, submitted_date")
        lines.append(") SELECT")
        lines.append(f"  (SELECT id FROM students ORDER BY id LIMIT 1 OFFSET {offset}),")
        lines.append(
            "  (SELECT p.id FROM professors p JOIN users u ON u.id = p.user_id"
            f" WHERE u.email = '{PROFESSOR_EMAIL}' LIMIT 1),"
        )
        for col in (
            "title",
            "technologies",
            "description",
            "category",
            "target_industry",
            "problem_statement",
            "proposed_solution",
            "unique_features",
            "innovation_aspect",
            "target_users",
            "expected_impact",
        ):
            lines.append(f"  '{esc_sql(p[col])}',")
        lines.append(f"  '{PROFESSOR_EMAIL}',")
        lines.append("  'PENDING',")
        lines.append("  CURRENT_DATE")
        lines.append("WHERE NOT EXISTS (")
        lines.append(f"  SELECT 1 FROM project_ideas WHERE title = '{esc_sql(p['title'])}'")
        lines.append(");")
        lines.append("")
    lines.extend(["COMMIT;", ""])
    return "\n".join(lines)


async def run_import(*, dry_run: bool = False) -> dict:
    ready_files = sorted(
        f.name
        for f in PROPOSALS_DIR.iterdir()
        if f.is_file()
        and f.suffix.lower() in (".pdf", ".docx")
        and f.name not in SKIP_FILENAMES
    )

    results: list[ImportResult] = []
    skipped_not_ready: list[str] = list(SKIP_FILENAMES)
    proposals: list[dict[str, str]] = []

    for filename in ready_files:
        try:
            record = build_proposal_record(filename)
            proposals.append(record)
        except Exception as exc:
            results.append(ImportResult(filename, "failed", reason=str(exc)))

    inserted_ids: list[dict] = []

    if not dry_run:
        async with AsyncSessionLocal() as db:
            prof = (
                await db.execute(
                    select(Professor)
                    .join(Professor.user)
                    .where(User.email == PROFESSOR_EMAIL)
                )
            ).scalar_one_or_none()
            if not prof:
                raise RuntimeError(f"{PROFESSOR_EMAIL} not found — run scripts.seed first")

            student_ids = [
                r[0]
                for r in (await db.execute(text("SELECT id FROM students ORDER BY id"))).all()
            ]
            if not student_ids:
                raise RuntimeError("No students found — run scripts.seed first")

            pre_count = await db.scalar(text("SELECT COUNT(*) FROM project_ideas"))

            for i, record in enumerate(proposals):
                existing = await db.execute(
                    select(ProjectIdea.id).where(ProjectIdea.title == record["title"])
                )
                if existing.scalar_one_or_none():
                    results.append(
                        ImportResult(
                            record["_source_file"],
                            "skipped_duplicate",
                            title=record["title"],
                            reason="Title already exists",
                        )
                    )
                    continue

                project = ProjectIdea(
                    student_id=student_ids[i % len(student_ids)],
                    professor_id=prof.id,
                    title=record["title"],
                    technologies=record["technologies"],
                    description=record["description"],
                    category=record["category"],
                    target_industry=record["target_industry"],
                    problem_statement=record["problem_statement"] or None,
                    proposed_solution=record["proposed_solution"] or None,
                    unique_features=record["unique_features"] or None,
                    innovation_aspect=record["innovation_aspect"] or None,
                    target_users=record["target_users"] or None,
                    expected_impact=record["expected_impact"] or None,
                    professor_email=PROFESSOR_EMAIL,
                    status=ProjectStatus.PENDING,
                )
                db.add(project)
                await db.flush()
                inserted_ids.append(
                    {"id": project.id, "title": project.title, "source": record["_source_file"]}
                )
                results.append(
                    ImportResult(
                        record["_source_file"],
                        "imported",
                        title=record["title"],
                        project_id=project.id,
                    )
                )

            await db.commit()
            post_count = await db.scalar(text("SELECT COUNT(*) FROM project_ideas"))
    else:
        pre_count = post_count = None
        for record in proposals:
            results.append(
                ImportResult(record["_source_file"], "dry_run", title=record["title"])
            )

    summary = {
        "total_found": len(list(PROPOSALS_DIR.glob("*"))),
        "ready_candidates": len(ready_files),
        "skipped_not_ready": skipped_not_ready,
        "proposals_built": len(proposals),
        "imported": sum(1 for r in results if r.status == "imported"),
        "skipped_duplicate": sum(1 for r in results if r.status == "skipped_duplicate"),
        "failed": sum(1 for r in results if r.status == "failed"),
        "dry_run": dry_run,
        "pre_count": pre_count,
        "post_count": post_count,
        "inserted": inserted_ids,
        "results": [r.__dict__ for r in results],
    }
    REPORT_JSON.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


def main() -> None:
    parser = argparse.ArgumentParser(description="Import university proposals into project_ideas")
    parser.add_argument("--dry-run", action="store_true", help="Parse only, no DB writes")
    parser.add_argument("--generate-sql", action="store_true", help="Write university_proposals_import.sql")
    parser.add_argument(
        "--import",
        dest="import_data",
        action="store_true",
        help="Insert into PostgreSQL (use with --generate-sql or alone)",
    )
    args = parser.parse_args()

    ready_files = sorted(
        f.name
        for f in PROPOSALS_DIR.iterdir()
        if f.is_file()
        and f.suffix.lower() in (".pdf", ".docx")
        and f.name not in SKIP_FILENAMES
    )
    proposals = [build_proposal_record(f) for f in ready_files]

    if args.generate_sql:
        SQL_OUT.write_text(generate_sql(proposals), encoding="utf-8")
        print(f"Wrote {SQL_OUT} ({len(proposals)} INSERT statements)")

    if args.dry_run:
        summary = asyncio.run(run_import(dry_run=True))
    elif args.generate_sql and not args.import_data:
        return
    else:
        summary = asyncio.run(run_import(dry_run=False))

    print(json.dumps({k: summary[k] for k in summary if k != "results"}, indent=2))


if __name__ == "__main__":
    main()
